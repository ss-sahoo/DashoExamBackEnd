"""
File Parser Service for extracting text from various file formats
Supports PDF extraction via Mathpix OCR for high-quality text extraction
"""
import os
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger('extraction')


class UnsupportedFileTypeError(Exception):
    """Raised when file type is not supported"""
    pass


class FileParsingError(Exception):
    """Raised when file parsing fails"""
    pass


class FileParserService:
    """Extract text content from various file formats"""
    
    SUPPORTED_MIME_TYPES = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
        'application/msword',  # .doc
        'text/plain',  # .txt
        'application/pdf',  # .pdf (via Mathpix)
        'image/jpeg',  # .jpg (via Mathpix)
        'image/png',  # .png (via Mathpix)
    ]
    
    SUPPORTED_EXTENSIONS = ['.txt', '.docx', '.doc', '.pdf', '.jpg', '.jpeg', '.png']
    
    def __init__(self):
        """Initialize the file parser service"""
        self.gemini_client = None  # Will be initialized when needed for image parsing
        self.mathpix_service = None  # Will be initialized when needed for PDF/image parsing
    
    def parse_file(self, file_path: str, file_type: str) -> str:
        """
        Extract text from file based on type
        
        Args:
            file_path: Path to uploaded file
            file_type: MIME type of file
            
        Returns:
            Extracted text content
            
        Raises:
            UnsupportedFileTypeError: If file type not supported
            FileParsingError: If parsing fails
        """
        if not os.path.exists(file_path):
            raise FileParsingError(f"File not found: {file_path}")
        
        # Validate file type
        if file_type not in self.SUPPORTED_MIME_TYPES:
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.SUPPORTED_EXTENSIONS:
                raise UnsupportedFileTypeError(
                    f"Unsupported file type: {file_type}. "
                    f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
                )
        
        try:
            # Route to appropriate parser based on file type
            if file_type == 'application/pdf' or file_path.endswith('.pdf'):
                return self.parse_pdf_mathpix(file_path)
            
            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword'] or file_path.endswith(('.docx', '.doc')):
                return self.parse_docx(file_path)
            
            elif file_type == 'text/plain' or file_path.endswith('.txt'):
                return self.parse_text(file_path)
            
            elif file_type in ['image/jpeg', 'image/png'] or file_path.endswith(('.jpg', '.jpeg', '.png')):
                return self.parse_image_mathpix(file_path)
            
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")
                
        except UnsupportedFileTypeError:
            raise
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}")
            raise FileParsingError(f"Failed to parse file: {str(e)}")
    
    def _get_mathpix_service(self):
        """Lazy initialization of Mathpix service"""
        if self.mathpix_service is None:
            try:
                from .mathpix_service import MathpixService
                self.mathpix_service = MathpixService()
            except Exception as e:
                logger.error(f"Failed to initialize Mathpix service: {e}")
                raise FileParsingError(f"Mathpix service not available: {str(e)}")
        return self.mathpix_service
    
    def parse_pdf_mathpix(self, file_path: str) -> str:
        """
        Extract text from PDF using Mathpix OCR API with caching.
        
        This provides high-quality OCR with:
        - LaTeX extraction for math equations
        - Table recognition
        - Support for scanned PDFs
        - Result caching to avoid redundant API calls
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content with LaTeX preserved
        """
        try:
            mathpix = self._get_mathpix_service()
            
            # Use cached extraction method (checks cache first, stores result after)
            text_content, ocr_result = mathpix.extract_pdf_with_cache(file_path)
            
            if not text_content or not text_content.strip():
                raise FileParsingError("No text content extracted from PDF")
            
            cache_status = "CACHED" if ocr_result.usage_count > 0 else "NEW"
            logger.info(
                f"Successfully extracted {len(text_content)} characters from PDF via Mathpix "
                f"[{cache_status}] (pages: {ocr_result.page_count})"
            )
            return text_content
            
        except FileParsingError:
            raise
        except Exception as e:
            logger.error(f"Mathpix PDF extraction failed: {str(e)}")
            # Fallback to PyPDF2 for simple text PDFs
            logger.info("Attempting fallback to PyPDF2...")
            return self.parse_pdf_pypdf2(file_path)
    
    def parse_image_mathpix(self, file_path: str) -> str:
        """
        Extract text from image using Mathpix OCR API with caching.
        
        Args:
            file_path: Path to image file
            
        Returns:
            Extracted text content with LaTeX preserved
        """
        try:
            mathpix = self._get_mathpix_service()
            
            # Use cached extraction method
            text_content, ocr_result = mathpix.extract_image_with_cache(file_path)
            
            if not text_content or not text_content.strip():
                raise FileParsingError("No text content extracted from image")
            
            cache_status = "CACHED" if ocr_result.usage_count > 0 else "NEW"
            logger.info(
                f"Successfully extracted {len(text_content)} characters from image via Mathpix "
                f"[{cache_status}]"
            )
            return text_content
            
        except Exception as e:
            logger.error(f"Mathpix image extraction failed: {str(e)}")
            raise FileParsingError(f"Failed to extract text from image: {str(e)}")
    
    def parse_pdf_pypdf2(self, file_path: str) -> str:
        """
        Fallback: Extract text from PDF using PyPDF2
        Used when Mathpix is unavailable or fails
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            import PyPDF2
            
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
            
            if not text_content:
                raise FileParsingError("No text content found in PDF")
            
            return '\n\n'.join(text_content)
            
        except ImportError:
            raise FileParsingError("PyPDF2 library not installed. Run: pip install PyPDF2")
        except Exception as e:
            raise FileParsingError(f"Failed to parse PDF: {str(e)}")
    
    def parse_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF - routes to Mathpix by default
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        return self.parse_pdf_mathpix(file_path)
    
    def _is_actually_text_file(self, file_path: str) -> bool:
        """
        Check if a file is actually plain text despite its extension
        
        DOCX files are ZIP archives and start with 'PK' (0x504B)
        If a file doesn't start with PK, it's likely plain text
        """
        try:
            with open(file_path, 'rb') as f:
                header = f.read(4)
            
            # DOCX/ZIP files start with PK (0x504B0304)
            if header[:2] == b'PK':
                return False
            
            # Check if content is readable as text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    f.read(1000)  # Try reading first 1000 chars
                return True
            except UnicodeDecodeError:
                return False
                
        except Exception:
            return False
    
    def parse_docx(self, file_path: str) -> str:
        """
        Extract text from Word document using python-docx
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        # First check if the file is actually plain text (common mistake)
        if self._is_actually_text_file(file_path):
            logger.warning(
                f"File {file_path} has .docx extension but is actually plain text. "
                "Parsing as text file instead."
            )
            return self.parse_text(file_path)
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            if not paragraphs:
                raise FileParsingError("No text content found in DOCX")
            
            return '\n\n'.join(paragraphs)
            
        except ImportError:
            raise FileParsingError("python-docx library not installed. Run: pip install python-docx")
        except Exception as e:
            # If DOCX parsing fails, try as plain text as fallback
            logger.warning(f"DOCX parsing failed, trying as plain text: {e}")
            try:
                return self.parse_text(file_path)
            except:
                raise FileParsingError(f"Failed to parse DOCX: {str(e)}")
    
    def parse_image(self, file_path: str) -> str:
        """
        Extract text from image using Gemini Vision API
        
        Args:
            file_path: Path to image file
            
        Returns:
            Extracted text content
            
        Note:
            This method uses Gemini's vision capabilities for OCR.
            The actual implementation will be in GeminiExtractionService
            to avoid circular dependencies.
        """
        # For now, return a placeholder
        # The actual image parsing will be handled by GeminiExtractionService
        # which has vision capabilities
        logger.info(f"Image file detected: {file_path}. Will use Gemini Vision API for extraction.")
        
        # Read image as bytes for later processing
        try:
            with open(file_path, 'rb') as f:
                image_data = f.read()
            
            # Return a marker that indicates this is an image file
            # The extraction service will handle it appropriately
            return f"[IMAGE_FILE:{file_path}]"
            
        except Exception as e:
            raise FileParsingError(f"Failed to read image file: {str(e)}")
    
    def parse_text(self, file_path: str) -> str:
        """
        Read plain text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content as string
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    
                    if content.strip():
                        return content
                    
                except UnicodeDecodeError:
                    continue
            
            raise FileParsingError("Could not decode text file with any supported encoding")
            
        except FileParsingError:
            raise
        except Exception as e:
            raise FileParsingError(f"Failed to read text file: {str(e)}")
    
    def validate_file_size(self, file_path: str, max_size_mb: int = 50) -> bool:
        """
        Validate file size
        
        Args:
            file_path: Path to file
            max_size_mb: Maximum allowed size in MB
            
        Returns:
            True if file size is valid
            
        Raises:
            FileParsingError: If file is too large
        """
        try:
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            
            if file_size > max_size_bytes:
                raise FileParsingError(
                    f"File size ({file_size / 1024 / 1024:.2f} MB) exceeds "
                    f"maximum allowed size ({max_size_mb} MB)"
                )
            
            return True
            
        except FileParsingError:
            raise
        except Exception as e:
            raise FileParsingError(f"Failed to check file size: {str(e)}")
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Get file information
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with file information
        """
        try:
            file_stat = os.stat(file_path)
            file_name = os.path.basename(file_path)
            file_ext = Path(file_path).suffix.lower()
            
            return {
                'name': file_name,
                'size': file_stat.st_size,
                'extension': file_ext,
                'path': file_path,
            }
            
        except Exception as e:
            logger.error(f"Failed to get file info: {str(e)}")
            return {}
